from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
CONTENT = ROOT / "INSTRUCTION_ADMINISTRATOR.md"
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "output"
QA = ROOT / "qa"
OUTPUT.mkdir(exist_ok=True)
QA.mkdir(exist_ok=True)


# compact_reference_guide tokens, with a named override for an easier-to-read
# operator document: 12 pt body text, 1.25 line spacing, and larger headings.
FONT = "Calibri"
INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(82, 91, 101)
MUTED = RGBColor(105, 113, 122)
LIGHT_FILL = "F4F6F9"
BLUE_FILL = "E8EEF5"
GOLD_FILL = "FFF8E8"
RED_FILL = "FFF1F1"
BORDER = "D7DBE2"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:%s" % edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in kwargs[edge]:
                element.set(qn("w:%s" % key), str(kwargs[edge][key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size: float, color: RGBColor = RGBColor(0, 0, 0), bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def remove_paragraph_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, 9, MUTED)


def prepare_image(path: Path) -> Path:
    prepared = QA / f"prepared_{path.stem}.jpg"
    if prepared.exists() and prepared.stat().st_mtime >= path.stat().st_mtime:
        return prepared
    with Image.open(path) as image:
        image = image.convert("RGB")
        image.thumbnail((1800, 1200), Image.Resampling.LANCZOS)
        image.save(prepared, "JPEG", quality=88, optimize=True)
    return prepared


def add_inline_runs(paragraph, text: str, size: float = 12, color: RGBColor = RGBColor(0, 0, 0)) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size, color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size, color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size, DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size, color)


def style_paragraph(paragraph, before=0, after=6, line=1.25, keep=False) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep
    remove_paragraph_border(paragraph)


def add_body(doc: Document, text: str, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=after)
    add_inline_runs(paragraph, text)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 18, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 14, BLUE, 14, 7
    else:
        size, color, before, after = 12.5, DARK_BLUE, 10, 5
    style_paragraph(paragraph, before=before, after=after, line=1.1, keep=True)
    run = paragraph.add_run(text)
    set_run_font(run, size, color, bold=True)


def new_numbering_id(doc: Document, abstract_num_id: int = 7) -> int:
    numbering = doc.part.numbering_part.element
    ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_id = max(ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_id


def add_list_item(doc: Document, text: str, ordered: bool = False, num_id: int | None = None) -> None:
    style = "List Number" if ordered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    if ordered and num_id is not None:
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num)
        p_pr.append(num_pr)
    add_inline_runs(paragraph, text)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": "2E74B5"}, top={"val": "single", "sz": "4", "color": BORDER}, bottom={"val": "single", "sz": "4", "color": BORDER}, right={"val": "single", "sz": "4", "color": BORDER})
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, before=0, after=0, line=1.2)
    add_inline_runs(paragraph, text.lstrip("> "), size=11.5, color=INK)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=4)


def add_placeholder(doc: Document, filename: str, caption: str) -> None:
    image_path = ASSETS / filename
    if image_path.exists():
        prepared = prepare_image(image_path)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(paragraph, before=4, after=4, line=1.0)
        run = paragraph.add_run()
        run.add_picture(str(prepared), width=Inches(6.35))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(cap, after=8, line=1.0)
        cap_run = cap.add_run(caption)
        set_run_font(cap_run, 9.5, MUTED, italic=True)
        return

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    set_cell_border(cell, top={"val": "single", "sz": "8", "color": BORDER}, bottom={"val": "single", "sz": "8", "color": BORDER}, left={"val": "single", "sz": "8", "color": BORDER}, right={"val": "single", "sz": "8", "color": BORDER})
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(46)
    paragraph.paragraph_format.space_after = Pt(46)
    run = paragraph.add_run(f"МЕСТО ДЛЯ СКРИНШОТА\n{filename}")
    set_run_font(run, 14, MUTED, bold=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(cap, after=8, line=1.0)
    cap_run = cap.add_run(caption + "  Добавьте файл в docs/admin_manual/assets и запустите builder снова.")
    set_run_font(cap_run, 9.5, MUTED, italic=True)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    widths = [TABLE_WIDTH_DXA // cols] * cols
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            style_paragraph(paragraph, after=0, line=1.12)
            add_inline_runs(paragraph, row[c_idx].strip() if c_idx < len(row) else "", size=10.5)
            if r_idx == 0:
                set_cell_shading(cell, BLUE_FILL)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, after=0, line=1.0)
    r = p.add_run("Motel  ·  Инструкция администратора")
    set_run_font(r, 9, MUTED, bold=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(kicker, after=12, line=1.0)
    kr = kicker.add_run("РАБОЧЕЕ РУКОВОДСТВО")
    set_run_font(kr, 11, BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(title, after=10, line=1.0)
    tr = title.add_run("Инструкция\nадминистратора")
    set_run_font(tr, 32, INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(subtitle, after=22, line=1.2)
    sr = subtitle.add_run("Motel: вход, смены, услуги, цены, оборудование и решение типовых проблем")
    set_run_font(sr, 14, GRAY)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [2700, 6660], indent=0)
    meta_rows = [("Версия", "1.0 · рабочая версия"), ("Для кого", "Администраторы и ответственные сотрудники"), ("Стиль", "Пошагово, простыми словами")]
    for idx, (label, value) in enumerate(meta_rows):
        for c, text in enumerate((label, value)):
            cell = meta.cell(idx, c)
            cell.text = ""
            set_cell_shading(cell, "F7F9FB" if idx % 2 == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.0)
            r = p.add_run(text)
            set_run_font(r, 10.5, INK if c == 0 else GRAY, bold=(c == 0))

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_placeholder(doc, "01_lock_screen.png", "Реальный экран блокировки из проекта. Пять быстрых нажатий по часам открывают вход администратора.")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(note, before=8, after=0, line=1.1)
    rr = note.add_run("Сначала прочитайте разделы 3–5. Они объясняют вход, выход и безопасную работу со сменой.")
    set_run_font(rr, 11, INK, bold=True)

    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "Содержание", 1)
    items = [
        "Что делает программа",
        "Перед началом работы",
        "Самый короткий путь: войти и выйти",
        "Что находится на экране «Настройки»",
        "Ежедневная работа со сменой",
        "Услуги, штрафы и цены",
        "Заставка и экран блокировки",
        "Оборудование и платежи",
        "Транзакции и диагностика",
        "Безопасность и сервер",
        "Telegram-уведомления",
        "Если что-то пошло не так",
        "Ежедневный чек-лист",
        "Что сфотографировать для полной версии инструкции",
        "Быстрый словарь кнопок",
    ]
    contents_num_id = new_numbering_id(doc)
    for item in items:
        add_list_item(doc, item, ordered=True, num_id=contents_num_id)
    add_callout(doc, "Совет: держите эту инструкцию рядом с терминалом, но не записывайте в нее пароли, токены и адрес сервера.")
    doc.add_page_break()


def parse_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    table_rows: list[list[str]] = []
    ordered_active = False
    ordered_num_id: int | None = None

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            flush_table()
            ordered_active = False
            ordered_num_id = None
            i += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r"[-: ]+", cell or " ") for cell in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        flush_table()
        if line.startswith("![SCREENSHOT:"):
            match = re.match(r"!\[SCREENSHOT:([^|]+)\|(.+)\]", line)
            if match:
                add_placeholder(doc, match.group(1).strip(), match.group(2).strip())
            i += 1
            continue
        if line.startswith("# "):
            # The title is already on the cover.
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
        elif line.startswith("> "):
            add_callout(doc, line)
        elif re.match(r"^\d+\.\s+", line):
            if not ordered_active:
                ordered_num_id = new_numbering_id(doc)
                ordered_active = True
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", line), ordered=True, num_id=ordered_num_id)
        elif line.startswith("- "):
            ordered_active = False
            ordered_num_id = None
            add_list_item(doc, line[2:].strip(), ordered=False)
        elif line.startswith("**") and line.endswith("**"):
            ordered_active = False
            ordered_num_id = None
            add_body(doc, line, after=6)
        else:
            ordered_active = False
            ordered_num_id = None
            add_body(doc, line)
        i += 1
    flush_table()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(12)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=0, line=1.0)
    r = p.add_run("Motel · версия 1.0 · страница ")
    set_run_font(r, 9, MUTED)
    add_page_field(p)


def build() -> Path:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_contents(doc)
    parse_markdown(doc, CONTENT.read_text(encoding="utf-8"))
    out = OUTPUT / "Motel_Инструкция_администратора.docx"
    doc.core_properties.title = "Инструкция администратора Motel"
    doc.core_properties.subject = "Пошаговое руководство по работе с терминалом Motel"
    doc.core_properties.author = "Motel"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build())
